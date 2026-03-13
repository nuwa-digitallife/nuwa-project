#!/usr/bin/env python3
"""
降临派手记 · 多Agent写作引擎

将交互式会话中验证过的写作流程沉淀为可复用的自动化代码。

架构：
  Pass 1: 写作Agent — 从素材涌现叙事，输出初稿
  Pass 2: 事实核查Agent — 穷举验证所有数字/日期/人名
  Pass 3: 审视Agent — Review checklist + 三层恰（只出报告，不改文章）
  Pass 3.5: 协商闭环 — Writing/Fact Agent 回应 → Review Agent 评估 → 达成共识 → 执行修改 → 验收
  Pass 5: 迭代求导（可选）— 证据硬度审计 → 定向调研 → 重写 → 收敛判断，循环
  Pass 4: 整合Agent — 合并修正，输出全部交付物
  Pass 4.5: 标题优化Agent — 标题备选 + 简介评估，注入 publish_guide

每个 Pass 是一次独立的 claude -p 调用，Pass 间通过文件通信。

使用：
  python engine.py --topic-dir wechat/公众号选题/2026-02-25|Anthropic蒸馏门 --persona 大史
  python engine.py --topic-dir ... --persona 章北海 --series 机器人系列
  python engine.py --topic-dir ... --persona 大史 --pass 2  # 只跑单个 pass（调试用）
  python engine.py --topic-dir ... --persona 大史 --iterate  # 启用迭代求导
  python engine.py --topic-dir ... --persona 大史 --pass 5 --iterate --max-iterations 1  # 只跑迭代
"""

import argparse
import logging
import os
import re
import signal
import subprocess
import sys
import time
from datetime import datetime
from pathlib import Path

# 添加 parent 到 path 以便直接运行
sys.path.insert(0, str(Path(__file__).parent))
from context_loader import ContextLoader
from image_collector import collect_images

# ── 配置 ────────────────────────────────────────────
PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent.parent  # nuwa-project/
PROMPTS_DIR = Path(__file__).resolve().parent / "prompts"
CONFIG_FILE = Path(__file__).resolve().parent.parent / "topic_config.yaml"

# 默认值（被 topic_config.yaml models.write_engine 覆盖）
DEFAULT_MODEL = "opus"
EFFORT_PER_PASS = {1: "high", 2: "high", 3: "high", 4: "medium"}
TOOLS_PER_PASS = {
    1: "Read,Grep,Glob,WebSearch,WebFetch",
    2: "WebSearch,WebFetch,Read",
    3: "WebSearch,WebFetch,Read",
    4: "Read",
}
TIMEOUT_PER_PASS = {1: 900, 2: 1200, 3: 1200, 4: 600}  # 秒；WebSearch 的 pass 给更长超时

# Pass 3.5 协商默认配置
CONSENSUS_CONFIG = {
    "max_rounds": 2,
    "num_revision_entries": 2,
}
EFFORT_PASS3B = {
    "write_respond": "high",
    "fact_respond": "high",
    "evaluate": "high",
    "revise": "high",
    "fact_revise": "high",
    "verify": "high",
}
TOOLS_PASS3B = {
    "write_respond": "Read",
    "fact_respond": "WebSearch,WebFetch,Read",
    "evaluate": "Read",
    "revise": "Read",
    "fact_revise": "WebSearch,WebFetch,Read",
    "verify": "Read",
}

# Pass 5 迭代求导默认配置
ITERATION_CONFIG = {"max_iterations": 2}
EFFORT_PASS5 = {
    "weakness": "high",
    "targeted_research": "high",
    "rewrite": "high",
    "compare": "medium",
}
TOOLS_PASS5 = {
    "weakness": "Read",
    "targeted_research": "WebSearch,WebFetch,Read",
    "rewrite": "Read",
    "compare": "Read",
}
TIMEOUT_PASS5 = {
    "weakness": 600,
    "targeted_research": 1200,
    "rewrite": 900,
    "compare": 600,
}

# Pass 4.5 标题优化默认配置
CONFIG_PASS4B = {
    "model": "sonnet",
    "effort": "medium",
    "tools": "Read",
    "timeout": 120,
}


def _load_model_config():
    """从 topic_config.yaml 加载配置，覆盖默认值。"""
    global DEFAULT_MODEL, EFFORT_PER_PASS, TOOLS_PER_PASS, TIMEOUT_PER_PASS
    global CONSENSUS_CONFIG, EFFORT_PASS3B, TOOLS_PASS3B
    global ITERATION_CONFIG, EFFORT_PASS5, TOOLS_PASS5, TIMEOUT_PASS5
    global CONFIG_PASS4B
    try:
        import yaml
        with open(CONFIG_FILE, "r", encoding="utf-8") as f:
            config = yaml.safe_load(f)
        we = config.get("models", {}).get("write_engine", {})
        if we.get("model"):
            DEFAULT_MODEL = we["model"]
        if we.get("pass_effort"):
            EFFORT_PER_PASS = {int(k): v for k, v in we["pass_effort"].items()}
        if we.get("pass_tools"):
            TOOLS_PER_PASS = {int(k): v for k, v in we["pass_tools"].items()}
        if we.get("pass_timeout"):
            TIMEOUT_PER_PASS = {int(k): v for k, v in we["pass_timeout"].items()}

        # Pass 3.5 协商配置
        if we.get("pass3b_effort"):
            EFFORT_PASS3B.update(we["pass3b_effort"])
        if we.get("pass3b_tools"):
            TOOLS_PASS3B.update(we["pass3b_tools"])

        # Pass 5 迭代求导配置
        if we.get("pass5_effort"):
            EFFORT_PASS5.update(we["pass5_effort"])
        if we.get("pass5_tools"):
            TOOLS_PASS5.update(we["pass5_tools"])
        if we.get("pass5_timeout"):
            TIMEOUT_PASS5.update(we["pass5_timeout"])

        # Pass 4.5 标题优化配置
        if we.get("pass4b"):
            CONFIG_PASS4B.update(we["pass4b"])

        # 共识循环参数
        consensus = config.get("consensus", {})
        if consensus:
            CONSENSUS_CONFIG.update(consensus)

        # 迭代参数
        iteration = config.get("iteration", {})
        if iteration:
            ITERATION_CONFIG.update(iteration)
    except Exception:
        pass  # 加载失败用默认值


_load_model_config()

logging.basicConfig(
    level=logging.INFO,
    format="[%(asctime)s] %(levelname)s: %(message)s",
    datefmt="%H:%M:%S",
)
log = logging.getLogger("write_engine")


# ── Prompt 模板加载 ──────────────────────────────────

def load_prompt_template(pass_name: str) -> str:
    """加载指定 Pass 的 prompt 模板。

    Args:
        pass_name: 模板名，如 "pass1_write", "pass3b_verify" 等
    """
    # 兼容旧接口：数字 → 文件名
    name_map = {
        "1": "pass1_write.md",
        "2": "pass2_factcheck.md",
        "3": "pass3_review.md",
        "4": "pass4_integrate.md",
    }
    if str(pass_name) in name_map:
        filename = name_map[str(pass_name)]
    else:
        filename = f"{pass_name}.md"

    path = PROMPTS_DIR / filename
    if not path.exists():
        log.error(f"Prompt 模板不存在: {path}")
        return ""

    template = path.read_text(encoding="utf-8")

    # 替换 {{SYSTEM_COMMON}}
    system_common_path = PROMPTS_DIR / "system_common.md"
    if system_common_path.exists():
        system_common = system_common_path.read_text(encoding="utf-8")
        template = template.replace("{{SYSTEM_COMMON}}", system_common)

    return template


def fill_template(template: str, context: dict) -> str:
    """将 {{KEY}} 占位符替换为实际内容。"""
    for key, value in context.items():
        template = template.replace(f"{{{{{key}}}}}", str(value))
    return template


# ── 限流重试助手 ──────────────────────────────────────

def _is_rate_limited(stderr: str, returncode: int = 0) -> bool:
    """检查是否应该重试（限流/过载/未知故障）。

    空 stderr + 非零退出码 → 很可能是限流（claude CLI 不总是输出错误信息）。
    """
    if returncode != 0 and not stderr.strip():
        return True  # 空 stderr + exit!=0 → 按限流处理，触发重试
    indicators = ["rate limit", "429", "overloaded", "too many requests",
                  "server error", "500", "502", "503"]
    stderr_lower = stderr.lower()
    return any(ind in stderr_lower for ind in indicators)


def run_claude_with_retry(cmd: list[str], prompt: str, timeout: int,
                          max_retries: int = 3, wait_seconds: int = 60,
                          cooldown_seconds: int = 300,
                          logger: logging.Logger = None) -> subprocess.CompletedProcess:
    """
    执行 claude -p 命令，遇到限流自动等待重试。

    策略：短间隔重试 max_retries 次（每次等 wait_seconds），
    全部失败后进入冷却（cooldown_seconds，默认5分钟），然后再试一轮。

    超时使用 Popen + 进程组杀掉，确保子进程也被终止。
    """
    _log = logger or logging.getLogger("write_engine")

    def _try_once():
        try:
            # 去掉 CLAUDECODE 环境变量，允许在 Claude Code 会话内嵌套调用
            env = {k: v for k, v in os.environ.items() if k != "CLAUDECODE"}
            # 使用 Popen + start_new_session 确保超时时能杀掉整个进程组
            proc = subprocess.Popen(
                cmd, stdin=subprocess.PIPE, stdout=subprocess.PIPE,
                stderr=subprocess.PIPE, text=True, encoding="utf-8",
                env=env, start_new_session=True,
            )
            try:
                stdout, stderr = proc.communicate(input=prompt, timeout=timeout)
                return subprocess.CompletedProcess(cmd, proc.returncode, stdout, stderr)
            except subprocess.TimeoutExpired:
                # 杀掉整个进程组（包括 claude 可能 spawn 的子进程）
                _log.error(f"claude -p 超时 ({timeout}s)，正在终止进程组...")
                try:
                    os.killpg(os.getpgid(proc.pid), signal.SIGTERM)
                    proc.wait(timeout=10)
                except (subprocess.TimeoutExpired, ProcessLookupError, OSError):
                    try:
                        os.killpg(os.getpgid(proc.pid), signal.SIGKILL)
                    except (ProcessLookupError, OSError):
                        pass
                    proc.wait(timeout=5)
                _log.error(f"进程已终止")
                return None
        except FileNotFoundError:
            _log.error("claude 命令不存在，请确认 Claude Code CLI 已安装")
            return None

    # 第一轮：短间隔重试
    for attempt in range(1, max_retries + 1):
        result = _try_once()
        if result is None:
            return None
        if result.returncode == 0 or not _is_rate_limited(result.stderr, result.returncode):
            return result
        if attempt < max_retries:
            _log.warning(f"限流等待 {wait_seconds}s... (重试 {attempt}/{max_retries})")
            time.sleep(wait_seconds)

    # 短间隔用尽，进入冷却
    minutes = cooldown_seconds / 60
    _log.warning(f"短间隔重试 {max_retries} 次用尽，冷却 {minutes:.0f}min...")
    time.sleep(cooldown_seconds)

    # 第二轮：冷却后再试一轮
    for attempt in range(1, max_retries + 1):
        result = _try_once()
        if result is None:
            return None
        if result.returncode == 0 or not _is_rate_limited(result.stderr, result.returncode):
            return result
        if attempt < max_retries:
            _log.warning(f"冷却后重试 {wait_seconds}s... (重试 {attempt}/{max_retries})")
            time.sleep(wait_seconds)

    _log.error("两轮重试均失败，放弃")
    return None


# ── Claude CLI 调用 ──────────────────────────────────

def run_claude(prompt: str, model: str, tools: str,
               effort: str = "high", timeout: int = 900) -> str:
    """调用 claude -p 执行单个 Pass。"""
    cmd = [
        "claude", "-p",
        "--model", model,
        "--allowedTools", tools,
        "--effort", effort,
        "--no-session-persistence",
    ]

    log.info(f"调用 claude -p (model={model}, effort={effort}, tools={tools})")
    log.info(f"Prompt 长度: {len(prompt)} chars ({len(prompt)//4} tokens approx)")

    result = run_claude_with_retry(cmd, prompt, timeout, logger=log)

    if result is None:
        return ""
    if result.returncode != 0:
        log.error(f"claude -p 失败 (exit={result.returncode})")
        log.error(f"stderr: {result.stderr[:500]}")
        return ""
    return result.stdout.strip()


# ── 输出解析 ─────────────────────────────────────────

def parse_delimited_output(output: str, delimiters: list[str]) -> dict:
    """解析用 ===DELIMITER=== 分隔的多段输出。"""
    result = {}
    current_key = None
    current_lines = []

    for line in output.split("\n"):
        matched = False
        for d in delimiters:
            if line.strip() == f"==={d}===" or line.strip() == f"```\n==={d}===":
                if current_key:
                    result[current_key] = "\n".join(current_lines).strip()
                current_key = d
                current_lines = []
                matched = True
                break
        if not matched and current_key:
            current_lines.append(line)

    if current_key:
        result[current_key] = "\n".join(current_lines).strip()

    return result


# ── 优化点类型解析 ───────────────────────────────────

def _has_type(review_report: str, type_marker: str) -> bool:
    """检查 review_report 中是否有指定类型的优化点。"""
    return type_marker in review_report


# ── Pass 1: 写作 ─────────────────────────────────────

def run_pass1(ctx: ContextLoader, topic_dir: Path, persona: str,
              series: str, model: str) -> str:
    """Pass 1: 写作Agent"""
    log.info("=" * 50)
    log.info("Pass 1: 写作Agent")
    log.info("=" * 50)

    template = load_prompt_template("1")
    context = ctx.assemble_pass1_context(topic_dir, persona, series)
    prompt = fill_template(template, context)

    output = run_claude(prompt, model, TOOLS_PER_PASS[1],
                        effort=EFFORT_PER_PASS[1], timeout=TIMEOUT_PER_PASS.get(1, 900))

    if output:
        out_path = topic_dir / "article_draft.md"
        out_path.write_text(output, encoding="utf-8")
        log.info(f"Pass 1 完成: {out_path} ({len(output)} chars)")
    else:
        log.error("Pass 1 失败: 无输出")

    return output


# ── Pass 2: 事实核查 ──────────────────────────────────

def run_pass2(ctx: ContextLoader, topic_dir: Path, article_draft: str,
              persona: str, model: str) -> tuple[str, str]:
    """Pass 2: 事实核查Agent。返回 (factcheck_report, corrected_article)。"""
    log.info("=" * 50)
    log.info("Pass 2: 事实核查Agent")
    log.info("=" * 50)

    template = load_prompt_template("2")
    context = ctx.assemble_pass2_context(topic_dir, article_draft, persona)
    prompt = fill_template(template, context)

    output = run_claude(prompt, model, TOOLS_PER_PASS[2],
                        effort=EFFORT_PER_PASS[2], timeout=TIMEOUT_PER_PASS.get(2, 1200))

    if not output:
        log.error("Pass 2 失败: 无输出")
        log.warning("⚠️ 事实核查被跳过，article_draft 中的 [待验证] 标记将保留到终稿")
        return "", article_draft

    parts = parse_delimited_output(output, ["FACTCHECK_REPORT", "CORRECTED_ARTICLE"])
    report = parts.get("FACTCHECK_REPORT", "")
    corrected = parts.get("CORRECTED_ARTICLE", "")

    if not corrected:
        log.warning("Pass 2 输出解析失败，使用原文继续")
        report = output
        corrected = article_draft

    if report:
        (topic_dir / "factcheck_report.md").write_text(report, encoding="utf-8")
        log.info(f"事实核查报告: {len(report)} chars")
    if corrected:
        (topic_dir / "article_factchecked.md").write_text(corrected, encoding="utf-8")
        log.info(f"修正后文章: {len(corrected)} chars")

    return report, corrected


# ── Pass 3: 审视（纯 Review，不改文章）──────────────

def run_pass3(ctx: ContextLoader, topic_dir: Path, article_factchecked: str,
              persona: str, series: str, model: str) -> str:
    """Pass 3: 审视Agent。只输出 review_report，不修改文章。"""
    log.info("=" * 50)
    log.info("Pass 3: 审视Agent（三层恰 Review — 只出报告）")
    log.info("=" * 50)

    template = load_prompt_template("3")
    context = ctx.assemble_pass3_context(topic_dir, article_factchecked, persona, series)
    prompt = fill_template(template, context)

    output = run_claude(prompt, model, TOOLS_PER_PASS[3],
                        effort=EFFORT_PER_PASS[3], timeout=TIMEOUT_PER_PASS.get(3, 1200))

    if not output:
        log.error("Pass 3 失败: 无输出")
        log.warning("⚠️ Review 被跳过，文章将不经过三层恰审视直接进入整合")
        return ""

    # Pass 3 只输出 REVIEW_REPORT
    parts = parse_delimited_output(output, ["REVIEW_REPORT"])
    report = parts.get("REVIEW_REPORT", "")

    if not report:
        # 如果没有分隔符，整段作为 report
        log.warning("Pass 3 输出没有 ===REVIEW_REPORT=== 分隔符，整段作为报告")
        report = output

    (topic_dir / "review_report.md").write_text(report, encoding="utf-8")
    log.info(f"Review 报告: {len(report)} chars")

    # 检测「改进行动」里是否有漏网的本篇建议
    _check_orphaned_recommendations(report, topic_dir)

    return report


# ── 改进行动漏网检测 ─────────────────────────────────

def _check_orphaned_recommendations(review_report: str, topic_dir: Path):
    """
    扫描 review_report 的「改进行动」/「本篇需修改」段落。
    如果里面有超出"已在优化点清单中详述"的实质内容，说明 Review Agent
    把本篇建议写到了优化点清单之外——这些内容不会进入 Pass 3.5 协商，会漏掉。

    检测到后：WARNING 报警 + 写到 orphaned_recommendations.md。
    """
    # 提取「改进行动」段落
    section_text = ""
    in_section = False
    for line in review_report.split("\n"):
        if "改进行动" in line:
            in_section = True
            continue
        if in_section:
            if line.startswith("## ") or line.startswith("```"):
                break
            section_text += line + "\n"

    if not section_text.strip():
        return

    # 提取「本篇需修改」部分（到「下篇需注意」之前）
    current_article_text = ""
    in_current = False
    for line in section_text.split("\n"):
        if "本篇需修改" in line or "本篇" in line:
            in_current = True
            current_article_text += line + "\n"
            continue
        if in_current:
            if "下篇" in line:
                break
            current_article_text += line + "\n"

    if not current_article_text.strip():
        return

    # 过滤掉模板默认文本
    filtered = current_article_text.strip()
    trivial_phrases = ["已在优化点清单中详述", "见上方优化点", "详见优化点"]
    for phrase in trivial_phrases:
        filtered = filtered.replace(phrase, "")

    # 去掉空行和纯标点行
    meaningful_lines = [
        l for l in filtered.split("\n")
        if l.strip() and len(l.strip()) > 5
    ]

    if not meaningful_lines:
        return

    # 发现漏网内容
    orphaned_content = "\n".join(meaningful_lines)
    log.warning("=" * 50)
    log.warning("⚠️ 发现漏网的本篇建议（不在优化点清单中，Pass 3.5 不会处理）：")
    for line in meaningful_lines:
        log.warning(f"  {line.strip()}")
    log.warning("已保存到 orphaned_recommendations.md，请人工检查是否需要补进优化点")
    log.warning("=" * 50)

    # 保存到文件
    out_path = topic_dir / "orphaned_recommendations.md"
    out_path.write_text(
        f"# 漏网的本篇建议\n\n"
        f"以下内容出现在 Review 报告的「改进行动/本篇需修改」段落中，\n"
        f"但不在结构化的优化点清单里。Pass 3.5 协商闭环不会处理这些内容。\n\n"
        f"**请人工检查是否需要补进优化点。**\n\n"
        f"---\n\n{orphaned_content}\n",
        encoding="utf-8",
    )


# ── CONSENSUS_DOC 精简助手 ────────────────────────────

def _compress_consensus_doc(consensus_doc: str, latest_evaluate: str) -> str:
    """将膨胀的 CONSENSUS_DOC 压缩为：未解决项 + 一段摘要。

    避免 Round 2+ 传入完整历史（10K→30K→50K 指数膨胀）。
    只保留：
    1. 未解决项列表（⏳ 标记的条目）
    2. 最后一轮评估的结论摘要（前 500 chars）
    """
    # 提取未解决项
    unresolved = []
    for line in consensus_doc.split("\n"):
        if "⏳" in line:
            unresolved.append(line.strip())

    # 取最后一轮评估的开头作为摘要
    summary = latest_evaluate[:500] if latest_evaluate else ""
    if len(latest_evaluate) > 500:
        summary += "\n[...评估摘要已截断...]"

    parts = ["# 前轮协商摘要（精简版）\n"]
    if unresolved:
        parts.append("## 未解决项\n")
        parts.extend(unresolved)
    if summary:
        parts.append(f"\n## 上轮评估摘要\n{summary}")

    return "\n".join(parts)


# ── Pass 3.5: 协商闭环 ──────────────────────────────

def run_consensus_loop(ctx: ContextLoader, topic_dir: Path,
                       review_report: str, article_factchecked: str,
                       persona: str, model: str,
                       skip_fact: bool = False,
                       max_rounds_override: int = None) -> tuple[str, str]:
    """
    Pass 3.5: 协商闭环。

    流程：
    1. Writing Agent 回应写作类优化点
    2. Fact Agent 回应事实类优化点（如有，且 skip_fact=False）
    3. Review Agent 评估所有回应，形成共识
    4. 如果还有未解决项且未达上限，回到 1
    5. Writing Agent 按共识执行写作类修改
    6. Fact Agent 按共识执行事实类修改（如有，且 skip_fact=False）
    7. Review Agent 验收

    Args:
        skip_fact: 如果 True，跳过 Fact Agent（当 Pass 2 已成功时，省 ~5-8K tokens）
        max_rounds_override: 覆盖 config 中的 max_rounds（CLI --consensus-rounds）

    Returns:
        (consensus_doc, final_article): 共识文档 和 验收后的文章
    """
    log.info("=" * 50)
    log.info("Pass 3.5: 协商闭环")
    log.info("=" * 50)

    max_rounds = max_rounds_override or CONSENSUS_CONFIG.get("max_rounds", 1)
    has_fact_issues = _has_type(review_report, "🔍") or _has_type(review_report, "🔀")
    has_write_issues = _has_type(review_report, "🖊️") or _has_type(review_report, "🔀")

    # Token 优化：如果 Pass 2 事实核查已完成，共识阶段跳过 Fact Agent
    if skip_fact and has_fact_issues:
        log.info("  Pass 2 事实核查已完成 → 共识阶段跳过 Fact Agent（省 token）")
        has_fact_issues = False

    if not has_write_issues and not has_fact_issues:
        log.info("Review 报告没有优化点，跳过协商")
        return "", article_factchecked

    consensus_doc = ""  # 累积协商记录
    consensus_doc_full = ""  # 完整版（保存用）
    latest_evaluate_text = ""  # 上轮评估文本（压缩用）

    # ── 协商轮次 ──
    for round_num in range(1, max_rounds + 1):
        log.info(f"── 协商第 {round_num}/{max_rounds} 轮 ──")

        # Token 优化：Round 2+ 使用精简版 consensus_doc
        if round_num > 1 and consensus_doc_full:
            consensus_doc = _compress_consensus_doc(consensus_doc_full, latest_evaluate_text)
            log.info(f"  共识文档精简: {len(consensus_doc_full)} → {len(consensus_doc)} chars")

        # Step 1: Writing Agent 回应
        if has_write_issues:
            log.info("  Writing Agent 回应写作类优化点...")
            template = load_prompt_template("pass3b_write_respond")
            context = ctx.assemble_write_respond_context(
                topic_dir, review_report, article_factchecked, consensus_doc, persona
            )
            prompt = fill_template(template, context)
            write_response = run_claude(
                prompt, model, TOOLS_PASS3B["write_respond"],
                effort=EFFORT_PASS3B["write_respond"]
            )
            parts = parse_delimited_output(write_response, ["WRITE_RESPONSE"])
            write_response_text = parts.get("WRITE_RESPONSE", write_response)
            round_entry = f"\n\n## 第 {round_num} 轮 — Writing Agent 回应\n\n{write_response_text}"
            consensus_doc += round_entry
            consensus_doc_full += round_entry
        else:
            log.info("  无写作类优化点，跳过 Writing Agent 回应")

        # Step 2: Fact Agent 回应
        if has_fact_issues:
            log.info("  Fact Agent 回应事实类优化点...")
            template = load_prompt_template("pass3b_fact_respond")
            context = ctx.assemble_fact_respond_context(
                topic_dir, review_report, article_factchecked, consensus_doc, persona
            )
            prompt = fill_template(template, context)
            fact_response = run_claude(
                prompt, model, TOOLS_PASS3B["fact_respond"],
                effort=EFFORT_PASS3B["fact_respond"]
            )
            parts = parse_delimited_output(fact_response, ["FACT_RESPONSE"])
            fact_response_text = parts.get("FACT_RESPONSE", fact_response)
            round_entry = f"\n\n## 第 {round_num} 轮 — Fact Agent 回应\n\n{fact_response_text}"
            consensus_doc += round_entry
            consensus_doc_full += round_entry
        else:
            log.info("  无事实类优化点，跳过 Fact Agent 回应")

        # Step 3: Review Agent 评估裁定
        log.info("  Review Agent 评估回应，形成共识...")
        template = load_prompt_template("pass3b_evaluate")
        context = ctx.assemble_consensus_evaluate_context(
            topic_dir, review_report, article_factchecked, consensus_doc, persona
        )
        prompt = fill_template(template, context)
        evaluate_output = run_claude(
            prompt, model, TOOLS_PASS3B["evaluate"],
            effort=EFFORT_PASS3B["evaluate"]
        )
        parts = parse_delimited_output(evaluate_output, ["CONSENSUS_UPDATE"])
        consensus_update = parts.get("CONSENSUS_UPDATE", evaluate_output)
        latest_evaluate_text = consensus_update
        round_entry = f"\n\n## 第 {round_num} 轮 — Review Agent 评估\n\n{consensus_update}"
        consensus_doc += round_entry
        consensus_doc_full += round_entry

        # 检查是否还有未解决项
        unresolved_count = consensus_update.count("⏳")
        log.info(f"  第 {round_num} 轮结果: {unresolved_count} 个未解决项")

        if unresolved_count == 0 or round_num >= max_rounds:
            if unresolved_count > 0:
                log.warning(f"  达到最大轮次 {max_rounds}，仍有 {unresolved_count} 个未解决项")
            break

    # ── 执行修改 ──
    log.info("── 执行共识修改 ──")

    current_article = article_factchecked
    all_change_lists = []

    # Step 4: Writing Agent 执行写作类修改
    if has_write_issues:
        log.info("  Writing Agent 执行写作类修改...")
        template = load_prompt_template("pass3b_revise")
        context = ctx.assemble_revision_context(
            topic_dir, current_article, consensus_doc, persona
        )
        prompt = fill_template(template, context)
        revise_output = run_claude(
            prompt, model, TOOLS_PASS3B["revise"],
            effort=EFFORT_PASS3B["revise"]
        )
        parts = parse_delimited_output(revise_output, ["REVISED_ARTICLE", "CHANGE_LIST"])
        revised = parts.get("REVISED_ARTICLE", "")
        change_list = parts.get("CHANGE_LIST", "")
        if revised:
            current_article = revised
            log.info(f"  写作类修改完成: {len(revised)} chars")
        else:
            log.warning("  Writing Agent 修改输出解析失败，使用当前文章继续")
        if change_list:
            all_change_lists.append(f"### Writing Agent 修改\n{change_list}")

    # Step 5: Fact Agent 执行事实类修改
    if has_fact_issues:
        log.info("  Fact Agent 执行事实类修改...")
        template = load_prompt_template("pass3b_fact_revise")
        context = ctx.assemble_fact_revision_context(
            topic_dir, current_article, consensus_doc, persona
        )
        prompt = fill_template(template, context)
        fact_revise_output = run_claude(
            prompt, model, TOOLS_PASS3B["fact_revise"],
            effort=EFFORT_PASS3B["fact_revise"]
        )
        parts = parse_delimited_output(fact_revise_output, ["REVISED_ARTICLE", "CHANGE_LIST"])
        revised = parts.get("REVISED_ARTICLE", "")
        change_list = parts.get("CHANGE_LIST", "")
        if revised:
            current_article = revised
            log.info(f"  事实类修改完成: {len(revised)} chars")
        else:
            log.warning("  Fact Agent 修改输出解析失败，使用当前文章继续")
        if change_list:
            all_change_lists.append(f"### Fact Agent 修改\n{change_list}")

    combined_change_list = "\n\n".join(all_change_lists)

    # Step 6: Review Agent 验收
    log.info("  Review Agent 验收...")
    template = load_prompt_template("pass3b_verify")
    context = ctx.assemble_verification_context(
        article_factchecked, current_article, consensus_doc,
        combined_change_list, topic_dir, persona
    )
    prompt = fill_template(template, context)
    verify_output = run_claude(
        prompt, model, TOOLS_PASS3B["verify"],
        effort=EFFORT_PASS3B["verify"]
    )
    parts = parse_delimited_output(verify_output, ["VERIFICATION", "VERIFIED_ARTICLE"])
    verification = parts.get("VERIFICATION", "")
    verified_article = parts.get("VERIFIED_ARTICLE", "")

    # 如果验收通过且有 VERIFIED_ARTICLE，用它；否则用 current_article
    final_article = verified_article if verified_article else current_article

    # 追加验收结果到共识文档（完整版）
    consensus_doc_full += f"\n\n## 验收结果\n\n{verification}"

    # 保存所有文件（用完整版 consensus_doc_full，便于溯源）
    (topic_dir / "consensus_doc.md").write_text(consensus_doc_full, encoding="utf-8")
    (topic_dir / "article_reviewed.md").write_text(final_article, encoding="utf-8")

    if verification:
        (topic_dir / "verification_report.md").write_text(verification, encoding="utf-8")

    log.info(f"Pass 3.5 完成: consensus_doc={len(consensus_doc_full)} chars, "
             f"final_article={len(final_article)} chars")

    return consensus_doc_full, final_article


# ── Pass 5: 迭代求导（证据硬化循环）──────────────────

def run_pass5_weakness(ctx: ContextLoader, topic_dir: Path, article: str,
                       version: int, persona: str, model: str) -> str:
    """Pass 5a: 证据硬度审计。返回 weakness_report。"""
    log.info(f"  5a: 证据硬度审计 (v{version})")

    template = load_prompt_template("pass5_weakness")
    context = ctx.assemble_pass5_weakness_context(topic_dir, article, persona)
    prompt = fill_template(template, context)

    output = run_claude(prompt, model, TOOLS_PASS5["weakness"],
                        effort=EFFORT_PASS5["weakness"],
                        timeout=TIMEOUT_PASS5.get("weakness", 600))

    if not output:
        log.error(f"  5a 失败: 无输出")
        return ""

    parts = parse_delimited_output(output, ["WEAKNESS_REPORT"])
    report = parts.get("WEAKNESS_REPORT", output)

    out_path = topic_dir / f"iteration_weakness_v{version}.md"
    out_path.write_text(report, encoding="utf-8")
    log.info(f"  5a 完成: {out_path.name} ({len(report)} chars)")

    return report


def run_pass5_targeted_research(ctx: ContextLoader, topic_dir: Path,
                                 article: str, weakness: str,
                                 version: int, persona: str, model: str) -> str:
    """Pass 5b: 定向调研。返回 targeted_research。"""
    log.info(f"  5b: 定向调研 (改进 v{version})")

    template = load_prompt_template("pass5_targeted_research")
    context = ctx.assemble_pass5_research_context(topic_dir, article, weakness, persona)
    prompt = fill_template(template, context)

    output = run_claude(prompt, model, TOOLS_PASS5["targeted_research"],
                        effort=EFFORT_PASS5["targeted_research"],
                        timeout=TIMEOUT_PASS5.get("targeted_research", 1200))

    if not output:
        log.error(f"  5b 失败: 无输出")
        return ""

    parts = parse_delimited_output(output, ["TARGETED_RESEARCH"])
    research = parts.get("TARGETED_RESEARCH", output)

    out_path = topic_dir / f"iteration_research_v{version}.md"
    out_path.write_text(research, encoding="utf-8")
    log.info(f"  5b 完成: {out_path.name} ({len(research)} chars)")

    return research


def run_pass5_rewrite(ctx: ContextLoader, topic_dir: Path,
                      article: str, weakness: str, research: str,
                      version: int, persona: str, model: str) -> str:
    """Pass 5c: 定向重写。返回 new_article。"""
    next_version = version + 1
    log.info(f"  5c: 定向重写 (v{version} → v{next_version})")

    template = load_prompt_template("pass5_rewrite")
    context = ctx.assemble_pass5_rewrite_context(
        topic_dir, article, weakness, research, persona
    )
    prompt = fill_template(template, context)

    output = run_claude(prompt, model, TOOLS_PASS5["rewrite"],
                        effort=EFFORT_PASS5["rewrite"],
                        timeout=TIMEOUT_PASS5.get("rewrite", 900))

    if not output:
        log.error(f"  5c 失败: 无输出")
        return ""

    parts = parse_delimited_output(output, ["REWRITTEN_ARTICLE"])
    new_article = parts.get("REWRITTEN_ARTICLE", output)

    out_path = topic_dir / f"article_v{next_version}.md"
    out_path.write_text(new_article, encoding="utf-8")
    log.info(f"  5c 完成: {out_path.name} ({len(new_article)} chars)")

    return new_article


def run_pass5_compare(ctx: ContextLoader, topic_dir: Path,
                      prev: str, curr: str,
                      v_prev: int, v_curr: int,
                      persona: str, model: str) -> tuple[str, bool]:
    """Pass 5d: 版本对比 + 收敛判断。返回 (comparison_report, converged)。"""
    log.info(f"  5d: 版本对比 (v{v_prev} ↔ v{v_curr})")

    template = load_prompt_template("pass5_compare")
    context = ctx.assemble_pass5_compare_context(topic_dir, prev, curr, persona)
    prompt = fill_template(template, context)

    output = run_claude(prompt, model, TOOLS_PASS5["compare"],
                        effort=EFFORT_PASS5["compare"],
                        timeout=TIMEOUT_PASS5.get("compare", 600))

    if not output:
        log.error(f"  5d 失败: 无输出")
        return "", False

    parts = parse_delimited_output(output, ["COMPARISON_REPORT", "VERDICT"])
    report = parts.get("COMPARISON_REPORT", output)
    verdict = parts.get("VERDICT", "").strip().upper()

    converged = "CONVERGED" in verdict and "NOT_CONVERGED" not in verdict

    # 累积追加到 iteration_comparison.md
    comparison_path = topic_dir / "iteration_comparison.md"
    header = f"\n\n## v{v_prev} → v{v_curr}\n\n"
    with open(comparison_path, "a", encoding="utf-8") as f:
        f.write(header + report + f"\n\n**Verdict**: {verdict}\n")

    log.info(f"  5d 完成: {'CONVERGED' if converged else 'NOT_CONVERGED'}")

    return report, converged


def run_iteration_loop(ctx: ContextLoader, topic_dir: Path,
                       article_v1: str, persona: str, model: str,
                       max_iterations: int = None) -> str:
    """
    Pass 5 迭代求导循环。

    v1 → 审计 → 调研 → 重写得到 v2 → 对比 → 收敛? → (回到审计)

    Returns: 最终版本的文章文本
    """
    if max_iterations is None:
        max_iterations = ITERATION_CONFIG.get("max_iterations", 2)

    log.info("=" * 50)
    log.info(f"Pass 5: 迭代求导（最多 {max_iterations} 轮）")
    log.info("=" * 50)

    current_article = article_v1
    current_version = 1

    # 保存 v1（Pass 4 产出）作为迭代起点备份
    v1_path = topic_dir / "article_v1.md"
    if not v1_path.exists():
        v1_path.write_text(article_v1, encoding="utf-8")
        log.info(f"  已保存迭代起点: article_v1.md ({len(article_v1)} chars)")

    # 清空旧的 iteration_comparison.md
    comparison_path = topic_dir / "iteration_comparison.md"
    comparison_path.write_text("# 迭代版本对比报告\n", encoding="utf-8")

    for iteration in range(1, max_iterations + 1):
        log.info(f"── 迭代第 {iteration}/{max_iterations} 轮 (当前 v{current_version}) ──")

        # 5a: 证据硬度审计
        weakness = run_pass5_weakness(
            ctx, topic_dir, current_article, current_version, persona, model
        )
        if not weakness:
            log.warning(f"  5a 失败，终止迭代")
            break

        # 检查是否全部是强/中
        if "VERDICT: ALL_STRONG" in weakness:
            log.info(f"  所有证据已足够硬，无需继续迭代")
            break

        # 5b: 定向调研
        research = run_pass5_targeted_research(
            ctx, topic_dir, current_article, weakness,
            current_version, persona, model
        )
        if not research:
            log.warning(f"  5b 失败，终止迭代")
            break

        # 5c: 定向重写
        new_article = run_pass5_rewrite(
            ctx, topic_dir, current_article, weakness, research,
            current_version, persona, model
        )
        if not new_article:
            log.warning(f"  5c 失败，终止迭代")
            break

        next_version = current_version + 1

        # 5d: 版本对比
        _, converged = run_pass5_compare(
            ctx, topic_dir, current_article, new_article,
            current_version, next_version, persona, model
        )

        current_article = new_article
        current_version = next_version

        if converged:
            log.info(f"  v{current_version} 已收敛，停止迭代")
            break

        if iteration < max_iterations:
            log.info(f"  v{current_version} 未收敛，继续下一轮")

    # 保存迭代最终版本（双写：交付物 + 溯源备份）
    (topic_dir / "article.md").write_text(current_article, encoding="utf-8")
    (topic_dir / "article_iterated.md").write_text(current_article, encoding="utf-8")
    log.info(f"Pass 5 完成: 最终版本 v{current_version} → article.md + article_iterated.md ({len(current_article)} chars)")

    return current_article


# ── Pass 4: 整合 ─────────────────────────────────────

def run_pass4(ctx: ContextLoader, topic_dir: Path, article_draft: str,
              factcheck_report: str, review_report: str,
              latest_article: str, consensus_doc: str,
              persona: str, model: str) -> bool:
    """Pass 4: 整合Agent。输出所有交付物。"""
    log.info("=" * 50)
    log.info("Pass 4: 整合Agent（交付物生成）")
    log.info("=" * 50)

    template = load_prompt_template("4")
    context = ctx.assemble_pass4_context(
        topic_dir, article_draft, factcheck_report, review_report,
        latest_article, consensus_doc, persona
    )
    prompt = fill_template(template, context)

    output = run_claude(prompt, model, TOOLS_PER_PASS[4],
                        effort=EFFORT_PER_PASS[4], timeout=TIMEOUT_PER_PASS.get(4, 600))

    if not output:
        log.error("Pass 4 失败: 无输出")
        return False

    parts = parse_delimited_output(output, [
        "ARTICLE", "ARTICLE_MDNICE", "POLL", "PUBLISH_GUIDE", "DESCRIPTION_OPTIONS"
    ])

    deliverables = {
        "article.md": parts.get("ARTICLE", ""),
        "article_mdnice.md": parts.get("ARTICLE_MDNICE", ""),
        "poll.md": parts.get("POLL", ""),
        "publish_guide.md": parts.get("PUBLISH_GUIDE", ""),
    }

    success = True
    for filename, content in deliverables.items():
        if content:
            (topic_dir / filename).write_text(content, encoding="utf-8")
            log.info(f"  {filename} ({len(content)} chars)")
        else:
            log.warning(f"  {filename} 为空")
            success = False

    desc = parts.get("DESCRIPTION_OPTIONS", "")
    if desc:
        (topic_dir / "description_options.md").write_text(desc, encoding="utf-8")

    return success


# ── Pass 4.5: 标题与元数据优化 ────────────────────────

def _update_publish_guide(topic_dir: Path, title: str, description: str):
    """将推荐标题和简介注入 publish_guide.md。"""
    pg_path = topic_dir / "publish_guide.md"
    if not pg_path.exists():
        log.warning("publish_guide.md 不存在，跳过注入")
        return

    content = pg_path.read_text(encoding="utf-8")

    # 替换或插入标题
    title_pattern = re.compile(r"\*\*标题[：:]\*\*.*")
    if title_pattern.search(content):
        content = title_pattern.sub(f"**标题：** {title}", content)
    else:
        # 在 # 发布指南 H1 后插入
        content = re.sub(
            r"(# 发布指南[^\n]*\n)",
            rf"\1\n**标题：** {title}\n",
            content,
            count=1,
        )

    # 替换或插入简介
    desc_pattern = re.compile(r"\*\*简介[：:]\*\*.*")
    if desc_pattern.search(content):
        content = desc_pattern.sub(f"**简介：** {description}", content)
    else:
        # 在标题行后插入
        content = re.sub(
            r"(\*\*标题[：:].*\n)",
            rf"\1**简介：** {description}\n",
            content,
            count=1,
        )

    pg_path.write_text(content, encoding="utf-8")
    log.info(f"  publish_guide.md 已更新标题和简介")


def run_pass4b(topic_dir: Path, force_title: str = None):
    """Pass 4.5: 标题与元数据优化。

    读取终稿 + 简介备选 + 当前标题，用 sonnet 生成优化标题和简介评估。
    如果 force_title 非空，直接写入 publish_guide 跳过 LLM。
    """
    log.info("=" * 50)
    log.info("Pass 4.5: 标题与元数据优化")
    log.info("=" * 50)

    # 读取终稿
    article_path = topic_dir / "article.md"
    if not article_path.exists():
        log.error("article.md 不存在，跳过 Pass 4.5")
        return
    article_text = article_path.read_text(encoding="utf-8")

    # 提取当前标题（H1）
    current_title = ""
    for line in article_text.split("\n"):
        if line.startswith("# "):
            current_title = line[2:].strip()
            break

    # 读取简介备选
    desc_path = topic_dir / "description_options.md"
    description_options = ""
    if desc_path.exists():
        description_options = desc_path.read_text(encoding="utf-8")

    # 强制标题模式：直接写入，跳过 LLM
    if force_title:
        log.info(f"使用强制标题: {force_title}")
        _update_publish_guide(topic_dir, force_title, "")
        return

    # 填充模板
    template = load_prompt_template("pass4b_title_meta")
    context = {
        "SHARED_CONTEXT": "",  # Pass 4.5 不需要共享上下文
        "ARTICLE_TEXT": article_text,
        "DESCRIPTION_OPTIONS": description_options or "（无简介备选）",
        "CURRENT_TITLE": current_title or "（无标题）",
    }
    prompt = fill_template(template, context)

    model = CONFIG_PASS4B["model"]
    effort = CONFIG_PASS4B["effort"]
    tools = CONFIG_PASS4B["tools"]
    timeout = CONFIG_PASS4B["timeout"]

    output = run_claude(prompt, model, tools, effort=effort, timeout=timeout)

    if not output:
        log.error("Pass 4.5 失败: 无输出")
        return

    # 解析输出
    parts = parse_delimited_output(
        output, ["TITLE_OPTIONS", "DESCRIPTION_EVALUATION", "RECOMMENDED"]
    )

    # 保存完整报告
    report_lines = [
        f"# 标题优化报告\n",
        f"**原标题：** {current_title}\n",
    ]
    if parts.get("TITLE_OPTIONS"):
        report_lines.append(f"## 备选标题\n\n{parts['TITLE_OPTIONS']}\n")
    if parts.get("DESCRIPTION_EVALUATION"):
        report_lines.append(f"## 简介评估\n\n{parts['DESCRIPTION_EVALUATION']}\n")
    if parts.get("RECOMMENDED"):
        report_lines.append(f"## 推荐\n\n{parts['RECOMMENDED']}\n")

    report = "\n".join(report_lines)
    (topic_dir / "title_options.md").write_text(report, encoding="utf-8")
    log.info(f"  title_options.md ({len(report)} chars)")

    # 提取推荐标题和简介
    recommended = parts.get("RECOMMENDED", "")
    rec_title = ""
    rec_desc = ""
    for line in recommended.split("\n"):
        m = re.match(r"\*\*标题[：:]\*\*\s*(.+)", line)
        if m:
            rec_title = m.group(1).strip()
        m = re.match(r"\*\*简介[：:]\*\*\s*(.+)", line)
        if m:
            rec_desc = m.group(1).strip()

    if rec_title or rec_desc:
        _update_publish_guide(topic_dir, rec_title or current_title, rec_desc)
        log.info(f"  推荐标题: {rec_title}")
        log.info(f"  推荐简介: {rec_desc[:50]}...")
    else:
        log.warning("Pass 4.5 输出解析失败，未能提取推荐标题/简介")


# ── 后处理：review→lessons 自动提取 ─────────────────

def extract_lessons_from_review(review_report: str, series: str, topic_dir: Path):
    """从 review_report 中只提取「下篇需注意」写入系列 lessons.md。

    注意：「本篇需修改」不提取——那些应该由 Pass 3.5 协商闭环处理。
    如果有漏网的本篇建议，_check_orphaned_recommendations 会报警。
    """
    if not series:
        log.info("独立篇，跳过 lessons 提取")
        return

    if not review_report:
        log.warning("review_report 为空，跳过 lessons 提取")
        return

    # 只提取「下篇需注意」段落
    lessons_text = ""
    in_next_article_section = False
    for line in review_report.split("\n"):
        if "下篇需注意" in line:
            in_next_article_section = True
            lessons_text += line + "\n"
            continue
        if in_next_article_section:
            # 遇到下一个 section 标题或分隔符就停
            if line.startswith("## ") or line.startswith("### 优化点") or line.startswith("```"):
                break
            lessons_text += line + "\n"

    if not lessons_text.strip():
        log.info("review_report 中未找到「下篇需注意」内容")
        return

    wechat_dir = PROJECT_ROOT / "wechat"
    lessons_path = wechat_dir / "公众号已发" / series / "lessons.md"
    lessons_path.parent.mkdir(parents=True, exist_ok=True)

    topic_name = topic_dir.name
    timestamp = datetime.now().strftime("%Y-%m-%d")
    entry = f"\n\n### 来自《{topic_name}》Review（{timestamp}，自动提取）\n\n{lessons_text.strip()}\n"

    with open(lessons_path, "a", encoding="utf-8") as f:
        f.write(entry)

    log.info(f"  lessons 已追加到: {lessons_path}")


# ── 后处理：封面图生成 ──────────────────────────────

def generate_cover_image(article_text: str, topic_dir: Path):
    """用 Gemini API 生成封面图。"""
    gen_image_path = Path(__file__).resolve().parent / "gen_image.py"
    if not gen_image_path.exists():
        log.warning("gen_image.py 不存在，跳过封面图生成")
        return

    lines = article_text.strip().split("\n")
    title = ""
    for line in lines[:20]:
        if line.startswith("# ") and not title:
            title = line[2:].strip()
            break

    cover_prompt = (
        f"Abstract digital art for a tech article cover image. "
        f"Topic: {title}. "
        f"Style: dark space background (#0a0e1a), subtle geometric patterns, "
        f"glowing cyan/teal accent lines, floating particles, futuristic tech aesthetic. "
        f"NO text, NO letters, NO words, NO characters. Pure visual/abstract design. "
        f"Aspect ratio: 900x383 (ultra-wide banner)."
    )

    output_path = topic_dir / "images" / "cover.png"
    output_path.parent.mkdir(parents=True, exist_ok=True)

    log.info("生成封面图 (Gemini)...")

    try:
        result = subprocess.run(
            [sys.executable, str(gen_image_path),
             "--prompt", cover_prompt,
             "--output", str(output_path),
             "--aspect", "900:383"],
            capture_output=True, text=True, timeout=120,
        )
        if result.returncode == 0:
            log.info(f"  封面图: {output_path}")
        else:
            log.warning(f"封面图生成失败: {result.stderr[:200]}")
    except subprocess.TimeoutExpired:
        log.warning("封面图生成超时 (120s)")
    except Exception as e:
        log.warning(f"封面图生成异常: {e}")


# ── 后处理：Word 文档生成 ──────────────────────────────

def generate_word_doc(article_text: str, topic_dir: Path):
    """将 article.md + 配图合并为 Word 文档，供用户审阅。"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        log.warning("python-docx 未安装，跳过 Word 生成 (pip install python-docx)")
        return

    doc = Document()

    # 基础样式
    style = doc.styles["Normal"]
    font = style.font
    font.name = "PingFang SC"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    images_dir = topic_dir / "images"

    # 封面图
    cover_path = images_dir / "cover.png"
    if cover_path.exists():
        try:
            doc.add_picture(str(cover_path), width=Inches(6))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            log.warning("  封面图格式不支持，跳过嵌入")

    # 逐行解析 Markdown → Word
    lines = article_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # 标题
        if line.startswith("# ") and not line.startswith("## "):
            p = doc.add_heading(line[2:].strip(), level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)

        # 图片标记 ![xxx](path)
        elif line.strip().startswith("!["):
            import re as _re
            m = _re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line.strip())
            if m:
                alt, img_path = m.group(1), m.group(2)
                # 尝试找本地图片
                full_path = None
                if (topic_dir / img_path).exists():
                    full_path = topic_dir / img_path
                elif (images_dir / Path(img_path).name).exists():
                    full_path = images_dir / Path(img_path).name
                # 按文件名模糊匹配
                if not full_path and images_dir.exists():
                    for f in images_dir.iterdir():
                        if f.suffix in (".png", ".jpg", ".jpeg") and f.name != "cover.png":
                            if alt and alt.lower() in f.name.lower():
                                full_path = f
                                break

                if full_path and full_path.exists():
                    try:
                        doc.add_picture(str(full_path), width=Inches(5.5))
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if alt:
                            cap = doc.add_paragraph(alt)
                            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            cap.runs[0].font.size = Pt(9)
                            cap.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                    except Exception:
                        p = doc.add_paragraph(f"[配图: {alt or img_path} (格式不支持)]")
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif not full_path:
                    # 图片不存在，写占位文字
                    p = doc.add_paragraph(f"[配图: {alt or img_path}]")
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.runs[0].font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

        # blockquote
        elif line.strip().startswith("> "):
            text = line.strip()[2:]
            p = doc.add_paragraph(text)
            p.paragraph_format.left_indent = Inches(0.5)
            p.runs[0].font.italic = True
            p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # 空行
        elif not line.strip():
            pass  # skip blank lines

        # 普通段落
        else:
            # 处理加粗和行内格式
            import re as _re
            clean = line.strip()
            if clean.startswith("- ") or clean.startswith("* "):
                clean = clean[2:]
                p = doc.add_paragraph(style="List Bullet")
            else:
                p = doc.add_paragraph()

            # 简单加粗解析
            parts = _re.split(r'(\*\*[^*]+\*\*)', clean)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

        i += 1

    # 保存
    output_path = topic_dir / "article_review.docx"
    doc.save(str(output_path))
    log.info(f"  Word 审阅文档: {output_path} ({output_path.stat().st_size // 1024}KB)")


# ── 主流程 ───────────────────────────────────────────

def run_engine(topic_dir: Path, persona: str, series: str = None,
               model: str = DEFAULT_MODEL, start_pass: int = 1,
               iterate: bool = False, max_iterations: int = None,
               consensus_rounds: int = None,
               skip_title: bool = False, force_title: str = None):
    """
    执行完整写作引擎流程。

    Pass 1 → Pass 2 → Pass 3 (纯 Review) → Pass 3.5 (协商闭环) → [Pass 5 迭代求导] → Pass 4 → Pass 4.5 (标题优化) → 后处理
    """
    topic_dir = Path(topic_dir).resolve()
    if not topic_dir.exists():
        log.error(f"选题目录不存在: {topic_dir}")
        return False

    log.info(f"写作引擎启动")
    log.info(f"  选题目录: {topic_dir}")
    log.info(f"  人设: {persona}")
    log.info(f"  系列: {series or '独立篇'}")
    log.info(f"  模型: {model}")
    log.info(f"  起始Pass: {start_pass}")
    _max_rounds = consensus_rounds or CONSENSUS_CONFIG.get("max_rounds", 1)
    log.info(f"  协商轮次上限: {_max_rounds}")
    if iterate:
        _max_iter = max_iterations or ITERATION_CONFIG.get("max_iterations", 2)
        log.info(f"  迭代求导: 开启 (最多 {_max_iter} 轮)")

    ctx = ContextLoader(PROJECT_ROOT)

    # ── Pass 1: 写作 ──
    if start_pass <= 1:
        article_draft = run_pass1(ctx, topic_dir, persona, series, model)
        if not article_draft:
            log.error("Pass 1 失败，终止")
            return False
    else:
        draft_path = topic_dir / "article_draft.md"
        final_path = topic_dir / "article.md"
        if draft_path.exists():
            article_draft = draft_path.read_text(encoding="utf-8")
            log.info(f"跳过 Pass 1，读取已有初稿")
        elif final_path.exists() and start_pass >= 5:
            article_draft = final_path.read_text(encoding="utf-8")
            log.info(f"跳过 Pass 1，使用 article.md 作为输入")
        else:
            log.error(f"Pass 1 被跳过但找不到初稿: {draft_path}")
            return False

    # ── Pass 2: 事实核查 ──
    if start_pass <= 2:
        factcheck_report, article_factchecked = run_pass2(
            ctx, topic_dir, article_draft, persona, model
        )
    else:
        fc_path = topic_dir / "article_factchecked.md"
        fr_path = topic_dir / "factcheck_report.md"
        article_factchecked = fc_path.read_text(encoding="utf-8") if fc_path.exists() else article_draft
        factcheck_report = fr_path.read_text(encoding="utf-8") if fr_path.exists() else ""
        log.info("跳过 Pass 2，读取已有核查结果")

    # ── Pass 3: 审视（纯 Review）──
    if start_pass <= 3:
        review_report = run_pass3(
            ctx, topic_dir, article_factchecked, persona, series, model
        )
    else:
        rr_path = topic_dir / "review_report.md"
        review_report = rr_path.read_text(encoding="utf-8") if rr_path.exists() else ""
        log.info("跳过 Pass 3，读取已有 Review 报告")

    # ── Pass 3.5: 协商闭环 ──
    consensus_doc = ""
    if start_pass <= 3:
        # Token 优化：如果 Pass 2 事实核查成功，共识阶段跳过 Fact Agent
        _skip_fact = bool(factcheck_report)
        consensus_doc, article_reviewed = run_consensus_loop(
            ctx, topic_dir, review_report, article_factchecked, persona, model,
            skip_fact=_skip_fact,
            max_rounds_override=consensus_rounds,
        )
    else:
        # 断点续跑时读取已有文件
        cd_path = topic_dir / "consensus_doc.md"
        ar_path = topic_dir / "article_reviewed.md"
        consensus_doc = cd_path.read_text(encoding="utf-8") if cd_path.exists() else ""
        article_reviewed = ar_path.read_text(encoding="utf-8") if ar_path.exists() else article_factchecked
        log.info("跳过 Pass 3.5，读取已有协商结果")

    # ── 降级状态检查 ──
    degraded_passes = []
    if not factcheck_report:
        degraded_passes.append("Pass 2 (事实核查)")
    if not review_report:
        degraded_passes.append("Pass 3 (审视Review)")
    if not consensus_doc:
        degraded_passes.append("Pass 3.5 (协商闭环)")
    if degraded_passes:
        log.warning("=" * 50)
        log.warning(f"⚠️ 降级模式：以下 Pass 被跳过 → {', '.join(degraded_passes)}")
        log.warning("终稿质量将低于完整流程，建议稳定网络后重跑失败的 Pass")
        log.warning("=" * 50)

    # ── Pass 5: 迭代求导（可选）──
    if iterate and start_pass <= 5:
        article_reviewed = run_iteration_loop(
            ctx, topic_dir, article_reviewed, persona, model,
            max_iterations=max_iterations,
        )
    elif start_pass == 5:
        # --pass 5 单独跑迭代
        article_reviewed = run_iteration_loop(
            ctx, topic_dir, article_reviewed, persona, model,
            max_iterations=max_iterations,
        )

    # ── Pass 4: 整合 ──
    if start_pass <= 4:
        success = run_pass4(
            ctx, topic_dir, article_draft, factcheck_report,
            review_report, article_reviewed, consensus_doc,
            persona, model
        )
    else:
        success = True

    # ── Pass 4.5: 标题与元数据优化 ──
    if success and not skip_title:
        try:
            run_pass4b(topic_dir, force_title=force_title)
        except Exception as e:
            log.warning(f"Pass 4.5 标题优化异常（不影响主流程）: {e}")

    # ── 后处理（全部 try-except 包裹，不影响主流程成功状态）──
    if success:
        try:
            extract_lessons_from_review(review_report, series, topic_dir)
        except Exception as e:
            log.warning(f"经验提取异常: {e}")

        final_article = ""
        try:
            if (topic_dir / "article.md").exists():
                final_article = (topic_dir / "article.md").read_text(encoding="utf-8")
            elif article_reviewed:
                final_article = article_reviewed
        except Exception as e:
            log.warning(f"读取终稿异常: {e}")

        if final_article:
            # 封面图（重试1次）
            for attempt in range(2):
                try:
                    generate_cover_image(final_article, topic_dir)
                    break
                except Exception as e:
                    log.warning(f"封面图生成异常 (attempt {attempt+1}): {e}")

            # 配图采集（重试1次）
            for attempt in range(2):
                try:
                    collect_images(final_article, topic_dir, model=model, effort="high")
                    break
                except Exception as e:
                    log.warning(f"配图采集异常 (attempt {attempt+1}): {e}")

            # Word 文档生成
            try:
                generate_word_doc(final_article, topic_dir)
            except Exception as e:
                log.warning(f"Word 文档生成异常: {e}")

        log.info("=" * 50)
        log.info("写作引擎完成")
        log.info("=" * 50)
        for f in sorted(topic_dir.iterdir()):
            if f.is_file() and f.suffix in (".md", ".docx"):
                log.info(f"  {f.name} ({f.stat().st_size} bytes)")
        images_dir = topic_dir / "images"
        if images_dir.exists():
            for f in sorted(images_dir.iterdir()):
                log.info(f"  images/{f.name} ({f.stat().st_size // 1024}KB)")
    else:
        log.warning("部分交付物缺失，请检查")

    return success


# ── CLI ──────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="降临派手记 · 多Agent写作引擎",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  # 完整流程（Pass 1→2→3→3.5→4）
  python engine.py --topic-dir wechat/公众号选题/2026-02-25|xxx --persona 大史

  # 指定系列（续恰用）
  python engine.py --topic-dir ... --persona 罗辑 --series 机器人系列

  # 从 Pass 3 断点续跑（跳过 1/2，从 Review 开始）
  python engine.py --topic-dir ... --persona 大史 --pass 3

  # 用 sonnet 跑快速测试
  python engine.py --topic-dir ... --persona 大史 --model sonnet
        """,
    )
    parser.add_argument("--topic-dir", required=True, help="选题目录路径")
    parser.add_argument("--persona", required=True, help="执笔人设（大史/章北海/罗辑）")
    parser.add_argument("--series", default=None, help="系列名（如 机器人系列）")
    parser.add_argument("--model", default=DEFAULT_MODEL, help=f"模型 (默认: {DEFAULT_MODEL})")
    parser.add_argument("--pass", type=int, default=1, dest="start_pass",
                        help="起始 Pass 编号 (1-5)，用于断点续跑")
    parser.add_argument("--iterate", action="store_true",
                        help="启用 Pass 5 迭代求导（证据硬化循环）")
    parser.add_argument("--max-iterations", type=int, default=None,
                        help="迭代最大轮次（默认从 topic_config.yaml 读取）")
    parser.add_argument("--consensus-rounds", type=int, default=None,
                        help="协商轮次上限（覆盖 topic_config.yaml，默认 1）")
    parser.add_argument("--skip-title", action="store_true",
                        help="跳过 Pass 4.5 标题优化")
    parser.add_argument("--title", default=None,
                        help="强制指定标题（跳过 LLM，直接写入 publish_guide）")

    args = parser.parse_args()
    success = run_engine(
        topic_dir=args.topic_dir,
        persona=args.persona,
        series=args.series,
        model=args.model,
        start_pass=args.start_pass,
        iterate=args.iterate,
        max_iterations=args.max_iterations,
        consensus_rounds=args.consensus_rounds,
        skip_title=args.skip_title,
        force_title=args.title,
    )
    sys.exit(0 if success else 1)


if __name__ == "__main__":
    main()
